from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
OUTPUT = DOCS / "PROPUESTA-COMERCIAL-DR-EDGAR-HERNANDEZ.docx"
LOGO = ROOT / "med-landing-dev" / "assets" / "images" / "brand" / "logo-horizontal-premium.png"

PRIMARY = "344729"
SECONDARY = "4A5942"
ACCENT = "7C2804"
GOLD = "C8B070"
SURFACE = "F7F7F7"
TEXT = "243126"
MUTED = "5F6B5A"
WHITE = "FFFFFF"
LIGHT_GOLD = "F3EBD5"
LIGHT_GREEN = "E9EEE6"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=11, bold=False, color=TEXT, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run(run, **kwargs)
    return run


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(TEXT)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, PRIMARY, 18, 10),
        ("Heading 2", 13, PRIMARY, 12, 6),
        ("Heading 3", 12, SECONDARY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Proposal Subtitle" not in [style.name for style in doc.styles]:
        subtitle = doc.styles.add_style("Proposal Subtitle", WD_STYLE_TYPE.PARAGRAPH)
        subtitle.font.name = "Calibri"
        subtitle.font.size = Pt(14)
        subtitle.font.color.rgb = rgb(SECONDARY)
        subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(8)


def add_header_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    add_text(paragraph, "DR. EDGAR E. HERNÁNDEZ  |  PROPUESTA WORDPRESS", size=8.5, bold=True, color=MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    add_text(paragraph, "Propuesta comercial · 8 de junio de 2026 · ", size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = p.add_run().add_picture(str(LOGO), width=Inches(3.9))
    logo._inline.docPr.set("descr", "Logo del Dr. Edgar E. Hernández, Nefrología")
    logo._inline.docPr.set("title", "Dr. Edgar E. Hernández")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(5)
    add_text(p, "PROPUESTA COMERCIAL", size=25, bold=True, color=PRIMARY)

    p = doc.add_paragraph(style="Proposal Subtitle")
    add_text(p, "Landing WordPress profesional y bilingüe", size=14, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    add_text(p, "Dr. Edgar E. Hernández · Nefrología", size=12, bold=True, color=SECONDARY)

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [1800, 7560])
    labels = ("Cobertura", "Fecha", "Preparado por", "Contacto")
    values = (
        "Xalapa y Boca del Río, Veracruz",
        "8 de junio de 2026",
        "Asael Marcial Grajales",
        "2281565984 · asael_marcial@hotmail.com",
    )
    for row, label, value in zip(table.rows, labels, values):
        set_cell_shading(row.cells[0], LIGHT_GREEN)
        p1 = row.cells[0].paragraphs[0]
        p2 = row.cells[1].paragraphs[0]
        add_text(p1, label, size=10, bold=True, color=PRIMARY)
        add_text(p2, value, size=10, color=TEXT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Honorario base", size=10, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    add_text(p, "MXN 6,000", size=27, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Total con IVA trasladado mediante CFDI: MXN 7,000", size=10.5, bold=True, color=PRIMARY)

    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(doc, text, bold_lead=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True)
        add_text(p, text[len(bold_lead):])
    else:
        add_text(p, text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        add_text(p, item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        add_text(p, item)


def add_callout(doc, title, body, fill=LIGHT_GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title, size=11, bold=True, color=PRIMARY)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_text(p, body, size=10.5, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_simple_table(doc, headers, rows, widths, header_fill=PRIMARY):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], header_fill)
        p = table.rows[0].cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, header, size=9.5, bold=True, color=WHITE)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            if len(table.rows) % 2 == 0:
                set_cell_shading(row.cells[index], SURFACE)
            p = row.cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index < len(row_values) - 1 else WD_ALIGN_PARAGRAPH.RIGHT
            add_text(p, value, size=9.5, bold=value.startswith("TOTAL"), color=TEXT)
        set_table_geometry(table, widths)
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    configure_styles(doc)
    add_header_footer(section)
    add_cover(doc)

    add_heading(doc, "1. Resumen de la propuesta")
    add_paragraph(
        doc,
        "Se propone diseñar, completar, configurar y publicar un sitio WordPress personalizado para el Dr. Edgar E. Hernández. El proyecto estará orientado a presentar sus servicios de nefrología, facilitar el contacto y la solicitud de citas, fortalecer su presencia profesional y apoyar el posicionamiento local en Xalapa y Boca del Río."
    )
    add_paragraph(
        doc,
        "El sistema utilizará un tema WordPress desarrollado específicamente para el proyecto, sin constructores visuales pesados. El cliente será propietario de su dominio, cuenta de hosting y contenidos. El costo de dominio, hosting, correo y servicios externos será contratado y pagado directamente por el cliente."
    )
    add_callout(
        doc,
        "Resultado esperado",
        "Un sitio médico profesional, rápido, bilingüe, administrable y preparado para captar solicitudes de cita y crecer mediante SEO local y contenidos de salud renal.",
    )

    add_heading(doc, "2. Qué incluirá el sitio")
    sections = [
        ("Identidad y experiencia visual", [
            "Identidad gráfica, logo, isotipo, favicon y paleta institucional.",
            "Diseño responsive para computadora, tableta y teléfono.",
            "Navegación, menú móvil, pie de página y llamadas a la acción.",
            "Optimización de imágenes y estructura ligera sin page builders.",
        ]),
        ("Páginas y contenido", [
            "Inicio, Sobre el Doctor, Servicios, Contacto y páginas para ambas sedes.",
            "Página individual para cada servicio médico aprobado.",
            "Blog de salud renal, preguntas frecuentes y testimonios autorizados.",
            "Páginas legales proporcionadas o aprobadas por el cliente.",
        ]),
        ("Contacto, mapas y conversión", [
            "Formulario de contacto y solicitud de cita mediante Fluent Forms.",
            "WhatsApp, click-to-call, mensajes predefinidos y notificaciones por correo.",
            "Mapas interactivos de Torre Hakim y Hospital MediMAC.",
            "Consentimiento de privacidad y botones para abrir rutas en Google Maps.",
        ]),
        ("SEO, bilingüe y crecimiento", [
            "Rank Math, sitemap, metadatos, canonical, robots y schema.",
            "SEO local para Xalapa y Boca del Río.",
            "Español e inglés mediante Polylang.",
            "Google Search Console, Analytics y Google Business cuando existan accesos.",
        ]),
    ]
    for heading, bullets in sections:
        add_heading(doc, heading, 2)
        add_bullets(doc, bullets)

    add_heading(doc, "3. Estado actual y trabajo restante")
    add_paragraph(
        doc,
        "Actualmente ya existe una base funcional del tema, la identidad visual, las páginas principales, las dos ubicaciones, los mapas y la estructura inicial de SEO."
    )
    add_numbered(doc, [
        "Recibir e integrar información profesional, clínica y fotografías aprobadas.",
        "Confirmar teléfono, WhatsApp, correo, horarios y contenido de servicios.",
        "Configurar formularios, privacidad, notificaciones y páginas legales.",
        "Configurar SEO, multilenguaje, blog, FAQ y testimonios autorizados.",
        "Ejecutar pruebas de accesibilidad, rendimiento, navegación y publicación.",
    ])

    add_heading(doc, "4. Información que deberá entregar el cliente")
    add_bullets(doc, [
        "Nombre profesional, cédulas y credenciales que deban publicarse.",
        "Biografía, formación, certificaciones y membresías verificables.",
        "Fotografía profesional y fotografías autorizadas.",
        "Teléfono, WhatsApp, correo, horarios y descripciones de servicios.",
        "Aviso de privacidad y textos legales revisados.",
        "Testimonios con autorización escrita, si se publicarán.",
        "Accesos o autorizaciones de dominio, hosting y servicios de Google.",
    ])
    add_callout(
        doc,
        "Criterio de publicación",
        "No se publicarán certificaciones, cifras, instituciones, testimonios ni afirmaciones médicas que no hayan sido confirmadas por el cliente.",
        fill=LIGHT_GREEN,
    )

    add_heading(doc, "5. Tiempo estimado")
    add_paragraph(
        doc,
        "El tiempo estimado para completar el proyecto es de 4 a 6 semanas."
    )

    add_heading(doc, "6. Inversión")
    add_simple_table(
        doc,
        ["Concepto", "Importe"],
        [
            ("Desarrollo completo de la landing web", "MXN 6,000.00"),
            ("IVA 16%, cuando corresponda emitir CFDI", "MXN 1,000.00"),
            ("TOTAL CON FACTURA", "MXN 7,000.00"),
            ("TOTAL SIN FACTURA", "MXN 6,000.00"),
        ],
        [6760, 2600],
    )
    add_paragraph(
        doc,
        "El precio incluye las Fases 1 y 2, versión bilingüe, configuración, pruebas y publicación."
    )

    add_heading(doc, "Forma de pago", 2)
    payment_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(payment_table, [4680, 4680])
    for index, heading in enumerate(("Sin solicitud de CFDI", "Con CFDI e IVA trasladado")):
        set_cell_shading(payment_table.rows[0].cells[index], PRIMARY if index == 0 else ACCENT)
        p = payment_table.rows[0].cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, heading, size=10, bold=True, color=WHITE)
    row = payment_table.add_row()
    set_table_geometry(payment_table, [4680, 4680])
    for index, lines in enumerate((
        ("50% al iniciar: MXN 3,000.00", "50% antes de publicar: MXN 3,000.00"),
        ("50% al iniciar: MXN 3,500.00", "50% antes de publicar: MXN 3,500.00"),
    )):
        set_cell_shading(row.cells[index], LIGHT_GREEN if index == 0 else LIGHT_GOLD)
        for line_index, line in enumerate(lines):
            p = row.cells[index].paragraphs[0] if line_index == 0 else row.cells[index].add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, line, size=10, bold=True, color=TEXT)

    add_callout(
        doc,
        "Nota fiscal",
        "El honorario comercial base es MXN 6,000.00. Cuando corresponda emitir CFDI, el total indicado en esta propuesta es MXN 7,000.00. Las retenciones u otras obligaciones deberán validarse con el contador antes de emitir el comprobante.",
    )

    add_heading(doc, "7. Hosting, dominio y WordPress")
    add_paragraph(
        doc,
        "WordPress.org no tiene costo de licencia. El dominio, hosting, correo y servicios externos no forman parte del honorario y serán contratados y pagados directamente por el cliente."
    )
    add_paragraph(
        doc,
        "Los precios fueron consultados el 8 de junio de 2026. Los importes en dólares usan una conversión ilustrativa de USD 1 = MXN 17.30. Las promociones, impuestos, tipo de cambio y renovaciones pueden variar."
    )

    hosting_rows = [
        ("Más barata a largo plazo", "Hostinger Premium", "MXN 1,679.52 / 48 meses", "MXN 419.88", "MXN 2,039.88"),
        ("Menor compromiso inicial", "DreamHost Launch", "USD 34.68 / 12 meses", "USD 34.68 · MXN 600", "USD 131.88 · MXN 2,282"),
        ("Costo-beneficio", "Hostinger Business", "MXN 2,831.52 / 48 meses", "MXN 707.88", "MXN 3,599.88"),
        ("Alternativa mexicana", "IONOS WordPress Start", "MXN 50/mes 6 meses + MXN 100/mes 6 meses", "MXN 900", "MXN 1,200"),
        ("Rendimiento y SEO", "SiteGround GrowBig", "USD 59.88 / 12 meses", "USD 59.88 · MXN 1,036", "USD 359.88 · MXN 6,226"),
        ("Referencia premium", "Kinsta Single 35k", "USD 350 anuales", "USD 350 · MXN 6,055", "USD 350"),
    ]
    add_simple_table(
        doc,
        ["Categoría", "Plan", "Pago inicial", "Promedio anual", "Renovación"],
        hosting_rows,
        [1800, 1750, 2350, 1850, 1610],
    )
    add_heading(doc, "Características principales", 2)
    add_bullets(doc, [
        "Hostinger Premium: dominio gratuito el primer año, SSL, correo, WordPress administrado y respaldos semanales.",
        "DreamHost Launch: dominio por un año, SSL, respaldos diarios, almacenamiento NVMe y correo gratuito durante los primeros tres meses.",
        "Hostinger Business: dominio inicial, SSL, correo, CDN, caché, respaldos diarios y staging.",
        "IONOS WordPress Start: pago y soporte en México, dominio .com o .mx, SSL, correo, respaldos diarios y actualizaciones administradas.",
        "SiteGround GrowBig: Google Cloud, CDN, caché multinivel, respaldos diarios, correo, staging y soporte WordPress.",
        "Kinsta: infraestructura premium y seguridad avanzada; no incluye dominio ni correo.",
    ])
    add_callout(
        doc,
        "Recomendación",
        "Hostinger Business es la recomendación general por costo-beneficio. Hostinger Premium reduce el promedio anual con cuatro años pagados por adelantado; DreamHost reduce el desembolso inicial; SiteGround prioriza rendimiento; Kinsta queda solo como referencia premium.",
        fill=LIGHT_GREEN,
    )
    add_paragraph(
        doc,
        "Ningún proveedor garantiza posicionamiento SEO. El beneficio procede de velocidad, disponibilidad, caché, CDN, seguridad y Core Web Vitals."
    )
    add_heading(doc, "Dominios", 2)
    add_bullets(doc, [
        "Usar .com como dominio principal.",
        "Registrar .com.mx adicionalmente para proteger la marca, si el cliente lo considera conveniente.",
        "Redirigir el dominio secundario al principal.",
        "Referencias de renovación: .com MXN 329.99 anuales y .com.mx MXN 612.99 anuales.",
        "IONOS permite elegir .com o .mx sin costo durante el primer año con contratación anual.",
    ])

    add_heading(doc, "Propuestas de nombre", 2)
    add_paragraph(
        doc,
        "La disponibilidad debe confirmarse con el registrador inmediatamente antes de contratar. Se priorizan nombres breves, fáciles de dictar y alineados con la marca profesional; incluir palabras clave en el dominio, por sí solo, no garantiza mejor posicionamiento."
    )
    add_simple_table(
        doc,
        ["Prioridad", "Dominio propuesto", "Motivo"],
        [
            ("Recomendada", "dredgarehernandez.com", "Coincide con la marca profesional y funciona fuera de México."),
            ("Alternativa corta", "dredgarhernandez.com", "Es fácil de recordar y escribir."),
            ("Enfoque México", "dredgarhernandez.mx", "Es breve y comunica presencia nacional."),
            ("Protección de marca", "dredgarhernandez.com.mx", "Refuerza la identificación con México y puede redirigir al .com."),
            ("Descriptiva", "nefrologoedgarhernandez.com", "Comunica nombre y especialidad, aunque es más larga."),
            ("Especialidad y apellido", "nefrologiahernandez.mx", "Es clara, relativamente corta y orientada al mercado mexicano."),
        ],
        [1900, 2850, 4610],
    )
    add_callout(
        doc,
        "Recomendación de dominio",
        "Registrar dredgarehernandez.com como dominio principal. Si el presupuesto lo permite, adquirir también la variante equivalente en .com.mx o .mx y redirigirla al dominio principal.",
        fill=LIGHT_GREEN,
    )

    add_heading(doc, "8. Referencia de mercado y valor del sitio propio")
    add_paragraph(
        doc,
        "Los siguientes precios fueron consultados el 8 de junio de 2026 en páginas oficiales. Son referencias comerciales y no comparaciones idénticas: Doctoralia incluye perfil, agenda y herramientas operativas, mientras DoctorWeb combina sitio, soporte y gestión de marketing."
    )
    add_simple_table(
        doc,
        ["Servicio publicado", "Precio mensual", "Referencia anual"],
        [
            ("Doctoralia Starter", "MXN 1,740 + IVA, facturado anualmente", "MXN 20,880 + IVA"),
            ("Doctoralia Plus", "MXN 2,340 + IVA, facturado anualmente", "MXN 28,080 + IVA"),
            ("Doctoralia VIP", "MXN 2,970 + IVA, facturado anualmente", "MXN 35,640 + IVA"),
            ("DoctorWeb Plan Google Despegue", "MXN 3,500 mensuales + MXN 2,000 iniciales de inversión", "MXN 44,000 durante el primer año"),
        ],
        [3200, 3100, 3060],
    )
    add_paragraph(
        doc,
        "DoctorWeb también publica planes de MXN 3,900 y MXN 5,900 mensuales, más importes iniciales. La inversión publicitaria puede variar según la especialidad y la estrategia seleccionada."
    )
    add_callout(
        doc,
        "Valor de esta propuesta",
        "El honorario de MXN 6,000 corresponde al desarrollo del sitio propio y no a una renta mensual. El cliente conserva el dominio, el sitio y sus contenidos; únicamente paga por separado hosting, dominio y servicios externos. Las plataformas anteriores pueden complementar el proyecto cuando se requieran agenda avanzada, publicidad administrada, CRM u otras funciones no incluidas.",
        fill=LIGHT_GOLD,
    )

    add_heading(doc, "9. Actualizaciones de información")
    add_paragraph(doc, "No se cobrará una mensualidad de mantenimiento.")
    add_paragraph(
        doc,
        "Se incluyen actualizaciones menores de información sin plazo definido y sin tiempo de respuesta garantizado, atendidas según disponibilidad."
    )
    add_bullets(doc, [
        "Correcciones de textos existentes.",
        "Actualización de teléfonos, correos, horarios y ubicaciones.",
        "Sustitución de imágenes existentes.",
        "Actualización de datos profesionales ya contemplados.",
    ])
    add_paragraph(
        doc,
        "Nuevas páginas, rediseños, funciones, integraciones, campañas, redacción extensa y recuperación por incidentes se cotizarán por separado."
    )
    add_paragraph(
        doc,
        "Las actualizaciones técnicas, seguridad y respaldos se apoyarán en el hosting administrado. Cualquier intervención técnica extraordinaria se evaluará por separado."
    )

    add_heading(doc, "10. Exclusiones y responsabilidades")
    add_heading(doc, "No incluido", 2)
    add_bullets(doc, [
        "Dominio, hosting, correo y licencias premium.",
        "Fotografía, video, publicidad pagada o administración de redes.",
        "Servicios legales, contables o regulatorios.",
        "Posicionamiento garantizado o promesas de resultados.",
        "Expediente clínico, cobros o agenda médica avanzada.",
    ])
    add_heading(doc, "Responsabilidad del cliente", 2)
    add_bullets(doc, [
        "Contratar y conservar dominio y hosting a su nombre.",
        "Renovar oportunamente servicios externos.",
        "Entregar información veraz y autorizada.",
        "Aprobar textos, fotografías, traducciones y contenido médico.",
        "Proporcionar avisos legales y accesos necesarios.",
    ])

    add_heading(doc, "11. Entrega y aceptación")
    add_paragraph(
        doc,
        "La entrega se considerará completa cuando el sitio esté instalado, funcionen sus páginas, navegación, formularios y enlaces, se haya integrado el contenido recibido, se haya revisado en móvil y escritorio y el cliente autorice la publicación."
    )
    add_paragraph(doc, "El segundo pago deberá liquidarse antes de publicar el sitio en el dominio definitivo.")

    add_heading(doc, "12. Fuentes consultadas")
    add_bullets(doc, [
        "Hostinger México — https://www.hostinger.com/mx/comprar-hosting",
        "Dominio .com Hostinger — https://www.hostinger.com/mx/tld/dominio-com",
        "Dominio .com.mx Hostinger — https://www.hostinger.com/mx/tld/dominio-com-mx",
        "DreamHost WordPress — https://www.dreamhost.com/es/wordpress/",
        "IONOS WordPress México — https://www.ionos.mx/alojamiento/wordpress-hosting",
        "SiteGround WordPress — https://www.siteground.com/es/wordpress-hosting.htm",
        "Kinsta WordPress — https://kinsta.com/es/precios/",
        "Doctoralia PRO, precios para especialistas — https://pro.doctoralia.com.mx/precios/medicos-y-especialistas",
        "DoctorWeb, planes de marketing médico — https://www.doctorweb.agency/planes.html",
        "Google Search, selección de un nombre de sitio — https://developers.google.com/search/docs/appearance/site-names",
        "Google Search, guía inicial de SEO — https://developers.google.com/search/docs/fundamentals/seo-starter-guide",
        "Google Search, sitios internacionales — https://developers.google.com/search/docs/specialty/international/managing-multi-regional-sites",
        "Google Search, Core Web Vitals — https://developers.google.com/search/docs/appearance/core-web-vitals",
        "SAT, Ley del IVA, artículo 1 — https://wwwmat.sat.gob.mx/articulo/19848/articulo-1",
    ])

    doc.add_page_break()
    add_heading(doc, "Aceptación")
    add_paragraph(doc, "Con la firma de este documento, ambas partes manifiestan conocer el alcance, inversión y condiciones generales de la propuesta.")
    signature = doc.add_table(rows=4, cols=2)
    set_table_geometry(signature, [4680, 4680])
    values = [
        ("CLIENTE", "PRESTADOR"),
        ("Dr. Edgar E. Hernández", "Asael Marcial Grajales"),
        ("Firma: __________________________", "Firma: __________________________"),
        ("Fecha: _________________________", "Fecha: _________________________"),
    ]
    for row_index, row_values in enumerate(values):
        for col_index, value in enumerate(row_values):
            if row_index == 0:
                set_cell_shading(signature.rows[row_index].cells[col_index], PRIMARY)
            p = signature.rows[row_index].cells[col_index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, value, size=10, bold=row_index in (0, 1), color=WHITE if row_index == 0 else TEXT)

    doc.core_properties.title = "Propuesta Comercial - Dr. Edgar E. Hernández"
    doc.core_properties.subject = "Landing WordPress profesional y bilingüe"
    doc.core_properties.author = "Asael Marcial Grajales"
    doc.core_properties.keywords = "WordPress, Nefrología, Propuesta comercial, Xalapa, Boca del Río"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
